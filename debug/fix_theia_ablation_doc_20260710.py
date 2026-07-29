from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(r"D:\daima\APT-Fusionstep2b1")
DOC_PATH = ROOT / "docs" / "tc3_path_reason_ablation_chain_windows_report_2026-07-09_zh.docx"
SUMMARY_PATH = ROOT / "docs" / "theia_ablation_chain_windows_summary_2026-07-10.json"
MONITOR_DIR = ROOT / "debug" / "remote_ops" / "out" / "theia_ablation_matrix_monitor_tmp"
WINDOW_DIR = MONITOR_DIR / "theia_window_files"
MARKER = "10. THEIA 消融实验（Window-Gated Baseline 与 2026-07-10 矩阵）"
NONE_TEXT = "\u65e0"
VARIANT_LABEL = {
    "baseline_current": "Baseline",
    "attack_prior_full": "\u542f\u7528\u653b\u51fb\u5148\u9a8c",
    "deterministic_mapping": "\u786e\u5b9a\u6027\u6620\u5c04",
    "no_claims_direct_mapping": "\u5173\u95ed Claims \u76f4\u6620\u5c04",
    "path_width_narrow": "\u7a84\u5bbd\u5ea6",
    "path_width_wide": "\u5bbd\u5bbd\u5ea6",
    "window_agg_top3": "\u7a97\u53e3\u805a\u5408 Top-3",
    "window_agg_top8": "\u7a97\u53e3\u805a\u5408 Top-8",
}
FILE_MAP = {
    "baseline_current": "baseline_tactic_diff_by_task.json",
    "attack_prior_full": "attack_prior_tactic_diff_by_task.json",
    "deterministic_mapping": "deterministic_tactic_diff_by_task.json",
    "no_claims_direct_mapping": "noclaims_tactic_diff_by_task.json",
    "path_width_narrow": "narrow_tactic_diff_by_task.json",
    "path_width_wide": "wide_tactic_diff_by_task.json",
    "window_agg_top3": "top3_tactic_diff_by_task.json",
    "window_agg_top8": "top8_tactic_diff_by_task.json",
}
WINDOW_ORDER = [
    "THEIA_20180410_1342_1342_02",
    "THEIA_20180410_1341_1455_01",
    "THEIA_20180412_1244_1326_03",
]


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8-sig"))


def build_summary_payload() -> dict:
    current = load_json(MONITOR_DIR / "current_comparison_20260710.json")
    old = load_json(MONITOR_DIR / "old_comparison_20260710.json")
    cur_results = {item["variant"]: item for item in current["results"]}
    old_results = {item["variant"]: item for item in old["results"]}

    summary_rows = []
    window_variant_rows = []
    for variant, zh_label in VARIANT_LABEL.items():
        cur_metrics = cur_results[variant]["metrics"]
        old_metrics = old_results[variant]["metrics"]
        for key in [
            "confirmed_window_recall",
            "strict_tactic_recall_macro",
            "strict_tactic_precision_macro",
            "off_window_high_risk_rate",
            "predicted_path_count",
            "predicted_path_with_report_count",
        ]:
            if cur_metrics[key] != old_metrics[key]:
                raise ValueError(f"current/c5d2981 mismatch on {variant}::{key}")

        tactic_rows = load_json(WINDOW_DIR / FILE_MAP[variant])
        by_window = {row["window_id"]: row for row in tactic_rows}
        recalls = []
        precisions = []
        for window_id in WINDOW_ORDER:
            row = by_window[window_id]
            gt = row["gt_tactics"]
            pred = row["predicted_tactics_union_top_n"]
            matched = row["matched_tactics"]
            missed = row["missed_tactics"]
            extra = row["extra_tactics"]
            recalls.append(len(matched) / len(gt) if gt else 0.0)
            precisions.append(len(matched) / len(pred) if pred else 0.0)
            window_variant_rows.append(
                {
                    "variant_key": variant,
                    "variant": zh_label,
                    "window_id": window_id,
                    "gt_tactics": ", ".join(gt),
                    "predicted_tactics": ", ".join(pred) if pred else NONE_TEXT,
                    "missed_tactics": ", ".join(missed) if missed else NONE_TEXT,
                    "extra_tactics": ", ".join(extra) if extra else NONE_TEXT,
                }
            )

        macro_recall = sum(recalls) / len(recalls)
        macro_precision = sum(precisions) / len(precisions)
        macro_f1 = 0.0
        if macro_recall + macro_precision > 0:
            macro_f1 = 2 * macro_recall * macro_precision / (macro_recall + macro_precision)
        path_count = cur_metrics["predicted_path_count"]
        report_count = cur_metrics["predicted_path_with_report_count"]
        summary_rows.append(
            {
                "variant_key": variant,
                "variant": zh_label,
                "status": cur_results[variant]["status"],
                "chain_window_count": len(WINDOW_ORDER),
                "hit_window_count": len(WINDOW_ORDER),
                "window_recall": cur_metrics["confirmed_window_recall"],
                "tactic_recall_macro": macro_recall,
                "tactic_precision_macro": macro_precision,
                "tactic_f1_macro": macro_f1,
                "predicted_path_count": path_count,
                "predicted_path_with_report_count": report_count,
                "report_rate": report_count / path_count if path_count else 0.0,
                "official_off_window_high_risk_rate": cur_metrics["off_window_high_risk_rate"],
            }
        )

    payload = {
        "code_versions_identical": True,
        "source_files": {
            "current_comparison": str(MONITOR_DIR / "current_comparison_20260710.json"),
            "old_comparison": str(MONITOR_DIR / "old_comparison_20260710.json"),
            "window_dir": str(WINDOW_DIR),
        },
        "chain_windows": WINDOW_ORDER,
        "summary_rows": summary_rows,
        "window_variant_rows": window_variant_rows,
    }
    SUMMARY_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return payload


def remove_existing_theia_section(doc: Document) -> None:
    body = doc._body._element
    marker_element = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("10. THEIA"):
            marker_element = paragraph._element
            break
    if marker_element is None:
        return

    children = list(body.iterchildren())
    start_index = children.index(marker_element)
    end_index = len(children)
    if children and children[-1].tag.endswith("sectPr"):
        end_index -= 1

    for child in children[start_index:end_index]:
        body.remove(child)


def add_summary_table(doc: Document, summary_rows: list[dict]) -> None:
    doc.add_paragraph("11. 子集整体结果")
    doc.add_paragraph("表 7  THEIA 链条窗口子集指标（current-code 与 c5d2981 结果一致）")
    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    headers = [
        "变体",
        "状态",
        "链条窗口数",
        "命中窗口数",
        "子集窗口召回",
        "子集战术召回(Macro)",
        "子集战术精度(Macro)",
        "子集战术F1(Macro)",
        "预测路径/有报告路径",
        "报告产出率",
    ]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text

    for row in summary_rows:
        cells = table.add_row().cells
        cells[0].text = row["variant"]
        cells[1].text = row["status"]
        cells[2].text = str(row["chain_window_count"])
        cells[3].text = str(row["hit_window_count"])
        cells[4].text = f'{row["window_recall"]:.3f}'
        cells[5].text = f'{row["tactic_recall_macro"]:.3f}'
        cells[6].text = f'{row["tactic_precision_macro"]:.3f}'
        cells[7].text = f'{row["tactic_f1_macro"]:.3f}'
        cells[8].text = f'{row["predicted_path_count"]} / {row["predicted_path_with_report_count"]}'
        cells[9].text = f'{row["report_rate"]:.3f}'


def add_window_table(doc: Document, window_rows: list[dict]) -> None:
    doc.add_paragraph("12. 窗口级战术结果")
    doc.add_paragraph("表 8  THEIA 有链条窗口的战术对照（采用 current-code 输出，c5d2981 完全一致）")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["变体", "窗口", "应有战术", "检测战术", "漏报", "误报"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text

    for row in window_rows:
        cells = table.add_row().cells
        cells[0].text = row["variant"]
        cells[1].text = row["window_id"]
        cells[2].text = row["gt_tactics"]
        cells[3].text = row["predicted_tactics"]
        cells[4].text = row["missed_tactics"]
        cells[5].text = row["extra_tactics"]


def add_analysis(doc: Document) -> None:
    doc.add_paragraph("13. 结果分析")
    doc.add_paragraph(
        "THEIA 上，Baseline、窄宽度、宽宽度三组结果完全一致：链条窗口子集的战术召回 Macro 均为 0.667，"
        "战术精度 Macro 均为 0.700，区别只体现在路径数量与报告产出率上。窄宽度把预测路径数从 126 降到 114，"
        "并把报告产出率从 0.786 提高到 0.868；宽宽度则把预测路径数抬到 130，同时把窗外高风险率从 0.160 推高到 0.173，"
        "但并未换来额外战术召回。"
    )
    doc.add_paragraph(
        "启用攻击先验是本轮最稳定的正收益：它把 THEIA 的链条窗口战术召回 Macro 从 0.667 提升到 0.722，"
        "同时把战术精度 Macro 从 0.700 提升到 0.783，且没有增加窗外高风险率。具体收益集中在 "
        "THEIA_20180410_1341_1455_01：该窗口在 Baseline 中只命中 COMMAND_AND_CONTROL 与 DISCOVERY，"
        "而攻击先验额外补回了 INITIAL_ACCESS，并去掉了多余的 CREDENTIAL_ACCESS 误报，只保留 EXFILTRATION 这一条额外战术。"
    )
    doc.add_paragraph(
        "确定性映射在召回上与攻击先验相同，也能把 THEIA_20180410_1341_1455_01 的 INITIAL_ACCESS 补回来，"
        "但它在两个窗口上引入了更多结构化误报：THEIA_20180410_1342_1342_02 多出 DISCOVERY，"
        "THEIA_20180410_1341_1455_01 多出 COLLECTION，因此精度 Macro 回落到 0.667。"
        "这说明在当前 THEIA 候选链质量下，最终性能提升并不只是前序链条决定，LLM 映射层仍然在抑制部分结构化过拟合。"
    )
    doc.add_paragraph(
        "关闭 Claims 直映射会让 3 个 confirmed 链条窗口全部失去战术输出，子集战术召回与精度都降为 0。"
        "结合它仍然保留 126 条候选路径和 99 条带报告路径这一事实，可以确认：Claims 不是单纯的解释文本包装，"
        "而是把候选链条转化为有效战术结论的核心承接层。"
    )
    doc.add_paragraph(
        "窗口聚合上限体现了另一个清晰折中。Top-3 把精度 Macro 提到全批最高的 0.867，但召回 Macro 降到 0.611，"
        "因为它在 THEIA_20180410_1341_1455_01 与 THEIA_20180412_1244_1326_03 上都收窄了可匹配路径，只保住更保守的 "
        "COMMAND_AND_CONTROL / DISCOVERY / INITIAL_ACCESS 子集。相反，Top-8 把召回 Macro 提到本批最高的 0.778："
        "它在 THEIA_20180410_1341_1455_01 中额外补回 EXECUTION 与 INITIAL_ACCESS，但也重新带回了 "
        "CREDENTIAL_ACCESS 与 EXFILTRATION 误报，因此精度回落到 0.756。"
    )
    doc.add_paragraph(
        "综合来看，THEIA 这组实验的最优折中是：若优先追求稳健综合表现，攻击先验优于其他变体；"
        "若强调保守输出，可考虑窗口聚合 Top-3；若优先追求更高召回，则窗口聚合 Top-8 最强，但需要接受更高的误报压力。"
        "无论 current-code 还是 c5d2981，这个结论都一致。"
    )


def main() -> None:
    summary = build_summary_payload()
    doc = Document(str(DOC_PATH))

    remove_existing_theia_section(doc)

    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph(MARKER)
    doc.add_paragraph(
        "本节补充 2026-07-10 完成的 THEIA 消融实验。统计口径继续限定为“只统计 confirmed 且真正产生候选攻击链条的窗口”。"
        "在本轮 THEIA 线上，这样的窗口共有 3 个：THEIA_20180410_1342_1342_02、THEIA_20180410_1341_1455_01、"
        "THEIA_20180412_1244_1326_03；attempted_failed 的 THEIA_20180413_1350_1404_04 不纳入本节指标。"
    )
    doc.add_paragraph(
        "本轮同时比较了 current-code 与 c5d2981 两套代码版本，其中 current-code 复用既有 THEIA baseline，"
        "c5d2981 从 fresh baseline 重新开始，并统一显式使用 +240 分钟偏移与 confirmed-only window gate。"
        "最终 8 个变体在两套代码上的窗口级与指标级结果完全一致，因此下表只保留一份合并结果。"
    )
    doc.add_paragraph(
        "消融因素包括：启用攻击先验、确定性映射、关闭 Claims 直映射、路径宽度收窄/放宽，以及窗口聚合上限从默认 Top-5 调为 "
        "Top-3 或 Top-8。由于这 3 个 confirmed 窗口都存在实际链条命中，所以本节中的“链条窗口子集指标”与 THEIA 本轮正式窗口指标相同。"
    )

    add_summary_table(doc, summary["summary_rows"])
    add_window_table(doc, summary["window_variant_rows"])
    add_analysis(doc)
    doc.save(str(DOC_PATH))
    print(DOC_PATH)


if __name__ == "__main__":
    main()
