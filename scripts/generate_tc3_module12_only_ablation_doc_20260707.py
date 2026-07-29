from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:\daima\APT-Fusionstep2b1")
SOURCE_JSON = (
    ROOT
    / "debug"
    / "remote_ops"
    / "out"
    / "tc3_module12_ablation_matrix_20260707_matrix_summary.remote.json"
)
OUT_DOCX = ROOT / "docs" / "tc3_module12_only_ablation_report_2026-07-07_zh.docx"


EXPERIMENT_META = {
    "late_fusion_on_current": {
        "group": "1.1/1.2/1.3 基线",
        "label": "当前方案：序列表示 + 图统计晚期融合（1:1），fanout>2，exclude segmented",
    },
    "late_fusion_off": {
        "group": "1.1 图统计晚期融合开关",
        "label": "关闭图统计晚期融合",
    },
    "early_concat_no_latefusion": {
        "group": "1.2 节点特征组织方式",
        "label": "早期拼接：序列表示 + 统计特征直接拼接后输入主干模型",
    },
    "sequence_only": {
        "group": "1.2 节点特征组织方式",
        "label": "仅使用序列表示，不使用统计特征",
    },
    "fanout_threshold2_include_segmented": {
        "group": "1.3 任务切图策略",
        "label": "fanout>2，include segmented children upstream",
    },
    "fanout_threshold3_exclude_segmented": {
        "group": "1.3 任务切图策略",
        "label": "fanout>3，exclude segmented children upstream",
    },
    "fanout_threshold3_include_segmented": {
        "group": "1.3 任务切图策略",
        "label": "fanout>3，include segmented children upstream",
    },
    "connected_components": {
        "group": "1.3 任务切图策略",
        "label": "connected components 整体连通切图",
    },
}


def load_data() -> dict:
    return json.loads(SOURCE_JSON.read_text(encoding="utf-8"))


def fmt_float(value: float | int | None, digits: int = 4) -> str:
    if value is None:
        return "-"
    if isinstance(value, int):
        return str(value)
    return f"{value:.{digits}f}"


def score_triplet(exp: dict) -> str:
    summary = exp.get("module2_fit_predict_summary", {})
    return (
        f"{fmt_float(summary.get('score_min'))} / "
        f"{fmt_float(summary.get('score_median'))} / "
        f"{fmt_float(summary.get('score_max'))}"
    )


def selected_tasks_text(exp: dict) -> str:
    items = exp.get("selected_positive_task_ids", [])
    if not items:
        return "-"
    return f"{len(items)}: " + ", ".join(items)


def set_run_font(run, size: int = 12, bold: bool = False, east_asia: str = "SimSun"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    style.font.size = Pt(12)


def set_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.2)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TRACE 与 CADETS 任务图检测消融实验报告")
    set_run_font(r, size=16, bold=True, east_asia="SimHei")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Module1/Module2-only Ablation Study for Task-Graph Detection")
    set_run_font(r, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026-07-07")
    set_run_font(r, size=10)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=14 if level == 1 else 12, bold=True, east_asia="SimHei")


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r, size=12)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(f"- {text}")
    set_run_font(r, size=11)


def set_cell(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=10, bold=True, east_asia="SimHei")


def add_design_table(doc: Document) -> None:
    add_table_caption(doc, "表 1  消融实验设计")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["编号", "变量", "对照设置", "研究目的"]
    for idx, text in enumerate(headers):
        set_cell(table.rows[0].cells[idx], text, bold=True, size=10)

    rows = [
        (
            "1.1",
            "图统计晚期融合开关",
            "当前方案 vs 关闭 late fusion",
            "验证图统计辅助分类器是否在 GraphSAGE 主干之外提供稳定增益",
        ),
        (
            "1.2",
            "节点特征组织方式",
            "late fusion / early concat / sequence only",
            "比较统计特征更适合作为独立分支，还是直接并入节点表示",
        ),
        (
            "1.3",
            "任务切图策略",
            "fanout 阈值、include/exclude segmented、connected",
            "分析切图粒度对任务数量、模块2指标和正类任务选取结果的影响",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value, size=9)


def add_result_table(doc: Document, dataset: str, rows: list[dict]) -> None:
    add_table_caption(doc, f"表 2{'A' if dataset == 'trace' else 'B'}  {dataset.upper()} 数据集消融结果")
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [
        "实验组",
        "具体设置",
        "任务图数",
        "Accuracy",
        "Precision",
        "Recall",
        "F1",
        "ROC-AUC",
        "正类任务与得分范围",
    ]
    for idx, text in enumerate(headers):
        set_cell(table.rows[0].cells[idx], text, bold=True, size=9)

    for exp in rows:
        meta = EXPERIMENT_META[exp["experiment_name"]]
        m1 = exp["module1_summary"]
        m2 = exp["module2_fit_predict_summary"]
        em = m2["evaluation_metrics"]
        detail = (
            f"{selected_tasks_text(exp)}\n"
            f"score(min/median/max)={score_triplet(exp)}"
        )
        values = [
            meta["group"],
            meta["label"],
            fmt_float(m1.get("task_count"), 0),
            fmt_float(em.get("accuracy")),
            fmt_float(em.get("precision")),
            fmt_float(em.get("recall")),
            fmt_float(em.get("f1")),
            fmt_float(em.get("roc_auc")),
            detail,
        ]
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell(cells[idx], value, size=8 if idx in (1, 8) else 9)


def trace_findings(doc: Document, trace_rows: list[dict]) -> None:
    add_heading(doc, "3 TRACE 实验结果分析", level=1)
    add_paragraph(
        doc,
        "在 TRACE 数据集上，当前 1:1 late fusion、早期拼接和仅序列表征三种特征组织方式，在模块2 fit_predict 指标上全部达到 1.0000，说明在当前任务图粒度下，节点特征组织方式并不是主要瓶颈。TRACE 的差异主要来自任务切图策略，而不是分类器结构本身。",
    )
    add_bullet(
        doc,
        "1.1 图统计晚期融合开关：开启或关闭 late fusion 都保持满分，说明 TRACE 上图统计分支不是决定性因素。",
    )
    add_bullet(
        doc,
        "1.2 节点特征组织方式：late fusion、early concat、sequence only 三者完全打平，说明序列表示本身已足够分离当前 TRACE 的正负任务图。",
    )
    add_bullet(
        doc,
        "1.3 任务切图策略：fanout>2, exclude segmented 维持当前稳定正类任务集合（task_0345/task_0546/task_0557/task_0558）；threshold 提高到 3 后，F1 直接降到约 0.498，说明切图过粗会显著破坏正类任务覆盖；connected 仅产出 14 张任务图，虽然表面指标满分，但只剩 task_0000 一个正类任务，失去任务粒度意义。",
    )


def cadets_findings(doc: Document, cadets_rows: list[dict]) -> None:
    add_heading(doc, "4 CADETS 实验结果分析", level=1)
    add_paragraph(
        doc,
        "CADETS 对图统计分支和切图粒度更敏感。与 TRACE 相比，CADETS 的任务图规模更大、结构更复杂，因此图统计 late fusion 和切图阈值都会显著改变模块2的召回、F1 以及最终选中的正类任务集合。",
    )
    add_bullet(
        doc,
        "1.1 图统计晚期融合开关：当前方案 F1=0.9281、Recall=0.9990、ROC-AUC≈1.0；关闭 late fusion 后 Recall 掉到 0.7500、F1 掉到 0.8326、ROC-AUC 掉到 0.6679，说明图统计辅助分支对 CADETS 明显提供稳定增益。",
    )
    add_bullet(
        doc,
        "1.2 节点特征组织方式：early concat 和 sequence only 与当前方案几乎打平，均保持 F1≈0.9281。这说明在当前配置下，CADETS 的主要收益来自“是否保留图统计辅助信息”，而不是统计特征到底早拼还是晚拼。",
    )
    add_bullet(
        doc,
        "1.3 任务切图策略：fanout>2 的 include/exclude 两种都接近当前基线；threshold 提高到 3 后，模块2指标反而提升到 1.0，但对应的任务图数量从 5247 缩到约 316x，且正类任务集合发生变化，因此这一结果更像“切图粒度变化带来的检测对象变化”，还不能直接推断为对下游攻击战术分析同样更优；connected 方案仅有 136 张任务图，F1 掉到 0.4909，是最差方案。",
    )


def add_conclusion(doc: Document) -> None:
    add_heading(doc, "5 结论与后续建议", level=1)
    add_paragraph(
        doc,
        "综合本轮 module1/module2-only 消融结果，可以得到三点结论。第一，TRACE 目前最稳的仍是 fanout>2 且 exclude segmented children upstream 的分支式切图策略；threshold=3 或 connected 都会显著改变甚至破坏正类任务粒度。第二，CADETS 上图统计晚期融合具有明确收益，不建议关闭。第三，CADETS 的 threshold=3 在模块2层面表现最好，但由于它显著改变了任务图数量和正类任务集合，因此需要单独进入后续 module3-module6 链路做进一步验证，不能仅凭模块2指标直接替换当前主线。",
    )
    add_bullet(doc, "若后续继续做任务图检测优化，建议优先保留 1:1 的 GraphSAGE/XGBoost 融合，并把切图策略作为主要研究变量。")
    add_bullet(doc, "若后续继续做攻击战术检测，应重点比较 CADETS 的 fanout>2 与 fanout>3 两类切图，在相同 GT 直筛口径下检查证据图、候选链和战术命中情况。")


def build_doc() -> None:
    data = load_data()
    doc = Document()
    set_normal_style(doc)
    set_page(doc)

    add_title(doc)

    add_heading(doc, "摘要", level=1)
    add_paragraph(
        doc,
        "本报告整理了 2026-07-07 在云端完成的 TRACE 与 CADETS 任务图检测消融实验。本轮实验只运行 module1 与 module2，统一采用“GraphSAGE 与 XGBoost 融合比率 1:1、数据增强倍数 count // 1000、先增强后划分训练集/测试集”的配置，并围绕图统计晚期融合开关、节点特征组织方式和任务切图策略三类变量开展对比。报告重点呈现任务图数量、模块2 fit_predict 指标、全量任务上的正类任务 ID，以及得分范围等任务图检测层面的结果，不包含下游攻击战术分析指标。",
    )

    add_heading(doc, "1 实验设置", level=1)
    add_bullet(doc, "实验对象：TRACE 与 CADETS 数据集。")
    add_bullet(doc, "运行范围：仅执行 module1 与 module2；每个实验都从 module1 重新开始跑。")
    add_bullet(doc, "统一配置：task_graph_stat_fusion_weight=0.5，task_tapas_augmentation_divisor=1000，task_tapas_augmentation_before_split=true。")
    add_bullet(doc, "模块2指标来源：fit_predict 的 holdout 评估；正类任务 ID 来源：同一训练模型下的 load_and_predict 全量预测。")
    add_design_table(doc)

    add_heading(doc, "2 结果总表", level=1)
    add_result_table(doc, "trace", data["results"]["trace"])
    doc.add_paragraph()
    add_result_table(doc, "cadets", data["results"]["cadets"])

    trace_findings(doc, data["results"]["trace"])
    cadets_findings(doc, data["results"]["cadets"])
    add_conclusion(doc)

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
