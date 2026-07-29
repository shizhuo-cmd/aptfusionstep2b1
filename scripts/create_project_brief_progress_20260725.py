from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"D:\daima\APT-Fusionstep2b1\docs\APT-Fusionstep2b1_项目简要进度说明_2026-07-25_zh_战术窗口明细版.docx")
BLUE = "2E74B5"
INK = "0B2545"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def font(run, size=11, bold=False, color="222222"):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_widths(table, widths):
    table.autofit = False
    pr = table._tbl.tblPr
    width = pr.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        pr.append(width)
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")
    layout = pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    ind = pr.first_child_found_in("w:tblInd")
    if ind is None:
        ind = OxmlElement("w:tblInd")
        pr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    for grid, value in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid.set(qn("w:w"), str(value))
    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(value))
            tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def cell_text(cell, text, bold=False, color="222222", align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    font(r, size=9.6, bold=bold, color=color)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    font(r, size=13, bold=True, color=BLUE)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    font(r, size=10.8)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    font(r, size=10.4)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    font(r, size=10.4)


def add_window_table(doc, rows):
    table = doc.add_table(rows=1, cols=5)
    table_widths(table, [1600, 2700, 2700, 1180, 1180])
    headers = ["攻击窗口", "应有战术", "检测战术", "漏报", "误报"]
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, LIGHT_BLUE)
        cell_text(cell, text, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, row_data)):
            cell_text(cell, text, bold=(index == 0), color=INK if index == 0 else "222222",
                      align=WD_ALIGN_PARAGRAPH.CENTER if index in (0, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)


def build():
    doc = Document()
    sec = doc.sections[0]
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, side, Inches(1))
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("项目简要进度说明"), size=8.5, color=MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("2026-07-25"), size=8.5, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run("APT-Fusionstep2b1 项目简要进度说明"), size=21, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    font(p.add_run("整理日期：2026-07-25"), size=10, color=MUTED)

    add_heading(doc, "目前整体进展")
    add_body(doc, "目前攻击战术分析部分已经比较完善。我们已经在 TRACE、CADETS 和 THEIA TC3 上反复按照攻击报告和原始日志核对过：先看攻击过程能不能被找出来，再看最后给出的战术是否合理。对那些已经有完整攻击过程的窗口，整体效果是比较好的。")
    add_body(doc, "不过，最近发现一个更前面的问题：之前使用的一些数据集里，切分出来的恶意任务图太少，而良性任务图数量非常多。这样会让训练和测试都很不平衡，也会限制后续攻击战术分析的上限。所以现在的重点逐渐转向寻找更适合做任务图检测和训练的新数据集。")

    add_heading(doc, "已有数据集上的攻击战术分析结果")
    add_body(doc, "下表只统计已经找到了完整攻击过程的攻击窗口，不把没有任务图或没有链条的空窗口算进去。这样反映的是后半段攻击战术分析本身的效果，而不是前面任务图是否落入时间窗口的问题。")
    table = doc.add_table(rows=1, cols=5)
    table_widths(table, [1450, 1750, 1750, 1750, 2660])
    headers = ["数据集", "有完整链条的窗口", "窗口命中率", "战术召回率", "战术精度"]
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, LIGHT_BLUE)
        cell_text(cell, text, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("TRACE", "2 个", "100%", "90.0%", "100.0%"),
        ("CADETS", "2 个", "100%", "83.3%", "100.0%"),
        ("THEIA TC3（step7b）", "2 个", "100%", "100.0%", "100.0%"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            cell_text(cell, text, bold=(cell is cells[0]), color=INK if cell is cells[0] else "222222",
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run("说明：THEIA TC3 另有一个攻击窗口没有匹配到任务图，因此没有纳入本表；它属于前面任务图和时间窗口对齐的问题，不应与战术分析本身混在一起计算。"), size=9.5, color=MUTED)

    add_heading(doc, "攻击战术分析是怎么做的")
    add_body(doc, "这里不是看到一个关键词就直接给出战术名称。整个过程大致分成下面四步，目的是让每个战术都能回到日志里的具体操作。")
    add_number(doc, "先从原始日志中恢复进程之间的父子关系，并把互相协作的一小组进程组织成任务图。这样可以把同一项活动里的进程、文件和网络操作放到一起看。")
    add_number(doc, "对进入分析范围的任务图，回查这些进程周围的文件访问、网络通信、执行、权限变化和清理操作，并保留时间顺序和对象关系。")
    add_number(doc, "把连续操作整理为可核对的攻击过程。例如，下载或生成可执行文件、修改权限后执行、读取敏感文件后外发、或者删除日志等，必须有对应的前后关系，不能只靠单个事件判断。")
    add_number(doc, "最后把已经有证据支持的攻击过程对照 ATT&CK 战术类别。评估时按攻击报告给出的时间窗口逐项比较：应有战术、检测战术、漏报和误报。")
    add_body(doc, "因此，表里的“检测战术”都应该能追到相应的进程、对象和时间顺序；如果前面没有形成完整攻击过程，就不把它算进这张战术分析表。")

    add_heading(doc, "有完整攻击链条窗口的逐项结果")
    add_body(doc, "以下使用中文战术名称，便于阅读：初始访问、执行、权限提升、命令与控制、发现、防御规避、收集、凭据访问。")
    add_body(doc, "TRACE")
    add_window_table(doc, [
        ("20180413\n12:43-12:53", "初始访问、执行、命令与控制、发现、防御规避", "初始访问、执行、命令与控制、发现、防御规避", "无", "无"),
        ("20180413\n13:50-14:28", "初始访问、执行、命令与控制、发现、收集", "初始访问、执行、命令与控制、发现", "收集", "无"),
    ])
    add_body(doc, "CADETS")
    add_window_table(doc, [
        ("20180412\n14:00-14:38", "初始访问、执行、权限提升、命令与控制、发现、防御规避", "初始访问、执行、权限提升、命令与控制", "发现、防御规避", "无"),
        ("20180413\n09:04-09:15", "初始访问、执行、权限提升、命令与控制、发现", "初始访问、执行、权限提升、命令与控制、发现", "无", "无"),
    ])
    add_body(doc, "THEIA TC3（step7b）")
    add_window_table(doc, [
        ("20180410\n13:42", "初始访问、执行、凭据访问", "初始访问、执行、凭据访问", "无", "无"),
        ("20180410\n13:41-14:55", "初始访问、执行、权限提升、命令与控制、发现、防御规避", "初始访问、执行、权限提升、命令与控制、发现、防御规避", "无", "无"),
    ])

    add_heading(doc, "为什么开始重点找新数据集")
    add_bullet(doc, "攻击战术分析要建立在“前面已经找到了比较完整的恶意任务图”这个前提上。")
    add_bullet(doc, "以前的一些数据集中，恶意任务图数量非常少，和大量良性任务图相比差距太大。模型很容易学成“全部判断为良性也很安全”，这会影响真正的检测能力。")
    add_bullet(doc, "目前看，TC3 的 THEIA 和 FiveDirections 两个数据集在任务图检测上更有希望：日志更适合处理，恶意任务图和良性任务图的数量关系也更接近实际训练需要。")

    add_heading(doc, "THEIA E5 的最新情况")
    add_body(doc, "我们已经试着接入了 THEIA E5 中一台主机的日志，也核对了恶意名单里记录的是进程、文件还是网络连接。文件和网络连接可以回查到对应进程，说明数据格式已经能够正确使用。")
    add_body(doc, "但这台主机仍然存在和旧数据集类似的问题：总共切出了 44,793 张任务图，其中能确认包含恶意进程的只有 18 张。正负样本差距仍然很大。另外，少数任务图特别大，里面混进了大量无关进程，会把少量恶意行为冲淡。")

    add_heading(doc, "后续计划")
    add_bullet(doc, "继续尝试 E5 里其他主机的日志，看是否能得到更合理的恶意和良性任务图比例。")
    add_bullet(doc, "继续尝试其他数据集，优先选择恶意样本数量更充足、日志字段更完整的数据。")
    add_bullet(doc, "对已经找到的大任务图做进一步拆分和核对，避免很多彼此无关的进程被放在同一张图里。")
    add_bullet(doc, "在新的数据集上先把任务图检测做稳定，再把现有的攻击战术分析流程接上去。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
