from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"D:\daima\APT-Fusionstep2b1")
RAW_JSON = ROOT / "docs" / "tc3_path_reason_ablation_step7bwidth_chain_windows_2026-07-09_raw.json"
OUT_DOCX = ROOT / "docs" / "tc3_path_reason_ablation_chain_windows_report_2026-07-09_zh.docx"


VARIANT_ORDER = [
    "baseline_current",
    "split_parent_inherit_off",
    "family_preserve_off_risk_topk",
    "attack_prior_full",
    "deterministic_mapping",
]


VARIANT_NAME = {
    "baseline_current": "Baseline",
    "split_parent_inherit_off": "2.1 关闭父标签继承",
    "family_preserve_off_risk_topk": "2.2 关闭关键攻击家族保留",
    "attack_prior_full": "2.3 启用攻击先验",
    "deterministic_mapping": "2.4 确定性战术映射",
}


DATASET_NAME = {
    "trace": "TRACE",
    "cadets": "CADETS",
}


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_page_border(section) -> None:
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        pg_borders.set(qn("w:offsetFrom"), "page")
        sect_pr.append(pg_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = pg_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            pg_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "24")
        element.set(qn("w:color"), "auto")


def fmt_float(value: float) -> str:
    return f"{value:.3f}"


def tactic_list(values) -> str:
    if not values:
        return "-"
    return ", ".join(values)


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_metric_table(doc: Document, dataset_key: str, dataset_info: dict) -> None:
    add_table_caption(doc, f"表 {1 if dataset_key == 'trace' else 2}  {DATASET_NAME[dataset_key]} 链条窗口子集指标")
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [
        "变体",
        "链条窗口数",
        "命中窗口数",
        "子集窗口召回",
        "子集战术召回(Macro)",
        "子集战术精度(Macro)",
        "预测路径/有报告路径",
    ]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, bold=True)

    variant_map = {v["variant"]: v for v in dataset_info["variants"]}
    for variant in VARIANT_ORDER:
        row = variant_map[variant]
        metrics = row["chain_window_metrics"]
        values = [
            VARIANT_NAME[variant],
            str(metrics["chain_window_count"]),
            str(metrics["chain_window_hit_count"]),
            fmt_float(metrics["chain_window_recall"]),
            fmt_float(metrics["chain_tactic_recall_macro"]),
            fmt_float(metrics["chain_tactic_precision_macro"]),
            f'{metrics["predicted_path_count"]}/{metrics["predicted_path_with_report_count"]}',
        ]
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            set_cell_text(cell, text)
    doc.add_paragraph("")


def add_window_table(doc: Document, dataset_key: str, dataset_info: dict) -> None:
    title = "表 3" if dataset_key == "trace" else "表 4"
    add_table_caption(doc, f"{title}  {DATASET_NAME[dataset_key]} 有链条窗口的战术对照")
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [
        "窗口",
        "变体",
        "最佳链条",
        "应有战术",
        "检测战术",
        "命中战术",
        "漏报战术",
        "误报战术",
    ]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, bold=True)

    variant_map = {v["variant"]: v for v in dataset_info["variants"]}
    for window_id in dataset_info["chain_window_ids"]:
        for variant in VARIANT_ORDER:
            row = variant_map[variant]
            window = next(x for x in row["window_rows"] if x["window_id"] == window_id)
            values = [
                window_id,
                VARIANT_NAME[variant],
                window["best_path_id"] or "-",
                tactic_list(window["gt_tactics"]),
                tactic_list(window["predicted_tactics"]),
                tactic_list(window["matched_tactics"]),
                tactic_list(window["missed_tactics"]),
                tactic_list(window["extra_tactics"]),
            ]
            cells = table.add_row().cells
            for cell, text in zip(cells, values):
                set_cell_text(cell, text)
    doc.add_paragraph("")


def add_paragraph(doc: Document, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10.5) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Pt(21) if align == WD_ALIGN_PARAGRAPH.JUSTIFY else Pt(0)


def build_document(data: dict) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    set_page_border(sec)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于 Step7b 宽度设定的攻击战术检测后半段消融实验报告")
    run.bold = True
    run.font.size = Pt(15)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRACE 与 CADETS：仅统计存在候选攻击链条的 confirmed 窗口")
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_paragraph(
        doc,
        "摘要：本文档整理 2026-07-08 在云端完成的后半段攻击战术检测消融实验。实验固定采用 "
        "module1_ground_truth_positive_base_only 入口、step7b 风格默认路径宽度、无数据增强、"
        "tactics-only 输出，并仅统计真正产生候选攻击链条的 confirmed 攻击窗口。消融因素包括："
        "分裂子任务父标签继承开关、关键攻击家族保留开关、攻击先验开关，以及大模型映射与确定性映射对比。"
        "结果表明：在当前链条质量下，关闭父标签继承或关闭关键攻击家族保留，对 TRACE 与 CADETS 的链条窗口子集指标几乎无影响；"
        "而启用攻击先验或切换到确定性映射，会导致两数据集的战术输出整体塌缩为零。",
    )

    add_paragraph(doc, "1. 实验设置", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
    add_paragraph(
        doc,
        "实验固定使用云端 step7bwidth module1-gt 路线。TRACE 仅保留具有时间匹配候选链条的两个 confirmed 窗口："
        "TRACE_20180413_1243_1253_04 与 TRACE_20180413_1350_1428_05；CADETS 仅保留两个 confirmed 链条窗口："
        "CADETS_20180412_1400_1438_03 与 CADETS_20180413_0904_0915_04。所有表格中的“指标”均为该窗口子集上的重新统计值，"
        "不再混入完全无链条的空窗。",
    )
    add_paragraph(
        doc,
        "四类消融的含义如下：2.1 关闭父标签继承，检验任务图切分后是否需要保留父任务根节点的状态/行为标签；"
        "2.2 关闭关键攻击家族保留，仅按纯风险分数执行 top-k；2.3 启用攻击先验，在候选路径和候选战术不变的前提下引入先验牵引；"
        "2.4 用确定性映射替代大模型映射，以区分最终性能来自前序链条质量还是映射层语义推理。",
    )

    add_paragraph(doc, "2. 子集整体结果", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
    add_metric_table(doc, "trace", data["datasets"]["trace"])
    add_metric_table(doc, "cadets", data["datasets"]["cadets"])

    add_paragraph(doc, "3. 窗口级战术结果", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
    add_window_table(doc, "trace", data["datasets"]["trace"])
    add_window_table(doc, "cadets", data["datasets"]["cadets"])

    add_paragraph(doc, "4. 结果分析", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
    add_paragraph(
        doc,
        "TRACE 上，Baseline、关闭父标签继承、关闭关键攻击家族保留三组结果完全一致：链条窗口召回均为 1.000，"
        "链条窗口战术召回 Macro 均为 0.900，差异只体现在路径数量上，其中关闭父标签继承将预测路径数从 44 降到 36，"
        "但并未改变两个链条窗口的战术结论。唯一稳定漏报来自 TRACE_20180413_1350_1428_05 的 COLLECTION。",
    )
    add_paragraph(
        doc,
        "CADETS 上，Baseline、关闭父标签继承、关闭关键攻击家族保留同样保持一致：链条窗口召回均为 1.000，"
        "链条窗口战术召回 Macro 为 0.833，稳定漏报集中在 CADETS_20180412_1400_1438_03 的 DISCOVERY 与 DEFENSE_EVASION。"
        "CADETS_20180413_0904_0915_04 在三组变体中均可完整命中应有战术。",
    )
    add_paragraph(
        doc,
        "与前述三组相比，启用攻击先验和切换到确定性映射会让 TRACE 与 CADETS 的 predicted_path_count 与 "
        "predicted_path_with_report_count 同时降为 0，链条窗口召回与战术召回也同步归零。这说明当前代码与配置下，"
        "攻击先验分支和确定性映射分支并未形成可比的稳定输出，若直接纳入主线会显著劣化结果。",
    )

    add_paragraph(doc, "5. 结论", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
    add_paragraph(
        doc,
        "从仅统计“有链条窗口”的结果看，2.1 与 2.2 这两类结构性开关暂未带来可观的窗口级收益，但也没有破坏当前基线；"
        "因此它们更像是影响搜索空间规模和解释结构的次级因素，而非当前性能瓶颈。真正的主要风险来自 2.3 与 2.4："
        "在现阶段链条与 claim 组织方式下，攻击先验和确定性映射都未能稳定承接前序输出，反而导致整窗无战术结果。"
        "后续若继续优化，应优先围绕 TRACE 的 COLLECTION 漏报与 CADETS 的 DISCOVERY/DEFENSE_EVASION 漏报进行路径重排、"
        "cleanup 尾段挂接或 claim 级规则修补，而不应先切换到攻击先验或确定性映射主线。",
    )

    return doc


def main() -> None:
    data = json.loads(RAW_JSON.read_text(encoding="utf-8"))
    doc = build_document(data)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
