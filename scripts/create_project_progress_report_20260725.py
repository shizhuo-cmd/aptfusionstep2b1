from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"D:\daima\APT-Fusionstep2b1\docs\APT-Fusionstep2b1_项目进度说明_2026-07-25_zh.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = table_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        table_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = table_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_text(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    # The style has no default text; apply explicit font for predictable Chinese rendering.
    set_run_font(r, size={1: 16, 2: 13, 3: 12}[level], bold=True,
                 color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=11, bold=True, color=INK)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=11, color="222222")
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color="222222")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=10.5, color="222222")
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=10.5, color="222222")
    return p


def add_callout(doc, label, text, color=INK):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(label + " ")
    set_run_font(r, size=10.5, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color="222222")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_status_table(doc):
    rows = [
        ("TRACE", "已完成多轮任务图与攻击行为识别回归。", "可用于观察规则变化是否伤到已有召回。", "仍有少量攻击窗口因上游任务图没有落入窗口而无链条。"),
        ("CADETS", "已完成任务图切分、线程处理、模型训练和攻击行为识别的多轮对照。", "已经确认超大根节点会把很多彼此独立的进程粘进同一张图。", "大图稀释恶意行为仍是主要问题，后续应优先处理根节点分支。"),
        ("THEIA TC3", "已完成日志字段核对、任务图边界诊断、时间对齐与后半段规则对照。", "文件路径解析和父子关系一致性问题已经定位并修正。", "不同版本任务图数量差异很大，需要固定同一份上游产物再比较。"),
        ("FiveDirections", "已完成原始日志接入、任务图检测和线程字段核查。", "数据量大、类别少，统计信息提取和训练划分已经单独验证。", "仍需要在稳定的训练划分和阈值选择上继续做对照。"),
        ("THEIA E5", "已完成 CDM20 日志接入、GT 格式核验、线程归并、对象关联和本地训练尝试。", "重新训练后的排序能力明显提升，但默认判定阈值仍没有选出测试集恶意图。", "正样本极少且存在超大任务图，是当前最需要解决的两件事。"),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_widths(table, [1200, 2750, 2900, 2510])
    headers = ["数据集", "目前做到哪", "已经确定的结论", "接下来的重点"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        add_text(cell, text, bold=True, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for idx, (cell, text) in enumerate(zip(cells, row_data)):
            add_text(cell, text, bold=(idx == 0), color=INK if idx == 0 else "222222", size=9.3,
                     align=WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_e5_table(doc):
    rows = [
        ("恶意名单总数", "70 条", "其中 64 条是进程，4 条是网络连接，2 条是文件。"),
        ("能直接对应到进程的恶意对象", "53 个", "日志里的线程合并后得到的唯一进程数。"),
        ("由文件或网络连接反查到的进程", "11 个", "它们都已经包含在前面的 53 个进程里，所以没有额外增加正样本。"),
        ("任务图总数", "44,793 张", "按当前规则从完整日志中切出的任务图。"),
        ("含已知恶意进程的任务图", "18 张", "训练数据极不平衡，是分类结果偏保守的根本原因。"),
        ("重新训练后的排序表现", "明显提高", "识别模型对恶意图的排序能力提升，但默认判定方式仍未挑出测试集正样本。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_widths(table, [2500, 1400, 5460])
    for cell, text in zip(table.rows[0].cells, ["检查项", "结果", "怎么理解"]):
        set_cell_shading(cell, LIGHT_BLUE)
        add_text(cell, text, bold=True, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for i, (cell, text) in enumerate(zip(cells, row_data)):
            add_text(cell, text, bold=(i == 0), color=INK if i == 0 else "222222", size=9.4,
                     align=WD_ALIGN_PARAGRAPH.CENTER if i == 1 else WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def build_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after, color in [
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, DARK_BLUE),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    # Header and footer.
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("APT-Fusionstep2b1 项目进度说明")
    set_run_font(header_run, size=8.5, color=MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("2026-07-25")
    set_run_font(footer_run, size=8.5, color=MUTED)

    # Opening block.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("项目进度说明")
    set_run_font(r, size=24, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("APT-Fusionstep2b1：从三个 TC3 数据集到 THEIA E5 的当前进展")
    set_run_font(r, size=13, color=MUTED)

    meta = doc.add_table(rows=3, cols=2)
    set_table_widths(meta, [1650, 7710])
    for (left, right), row in zip([
        ("整理日期", "2026-07-25"),
        ("阅读对象", "希望快速了解项目进展、已证实问题和下一步方向的同学"),
        ("写法", "尽量用日常语言描述；代码里的内部名称和缩写不作为正文概念使用"),
    ], meta.rows):
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        add_text(row.cells[0], left, bold=True, color=INK, size=10)
        add_text(row.cells[1], right, color="222222", size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    add_callout(
        doc,
        "一句话进展：",
        "三个 TC3 数据集的整条检测流程已经做过多轮对照，当前最清楚的结论是：结果好不好，不只取决于最后的战术判断，还非常依赖前面任务图切得是否合理、恶意名单是否对齐、以及证据是否被完整保留下来。最新的 THEIA E5 已经成功接入，但它的正样本太少、部分任务图太大，当前优先要先把这两个上游问题处理好。",
    )

    add_heading(doc, "1. 这个项目现在在做什么", 1)
    add_body(doc, "项目目标是从系统日志里找出可疑进程及其相关操作，再把这些操作串成攻击过程，最后给出攻击者可能采取的战术。为了避免只看一个分数而误判，我们一直把流程拆开检查：先看日志和已知恶意名单能不能对上，再看任务图是否把相关进程放到一起，随后看攻击过程是否被保住，最后才看战术名称是否判断正确。")
    add_body(doc, "这段时间的工作重点有两条：一条是把 TRACE、CADETS、THEIA 这三个 TC3 数据集的评估做扎实，另一条是把新的 THEIA E5 数据集接进来，并验证它的日志格式、恶意名单格式和训练方法是否真的适用。")

    add_heading(doc, "2. 到目前为止，已经做了哪些工作", 1)
    add_heading(doc, "2.1 三个 TC3 数据集：先把“该看什么”对齐", 2)
    add_body(doc, "我们重新整理了攻击报告中的已确认攻击行为、未成功但有明确攻击动作的尝试，以及每个攻击时间段应当出现的战术。评估时不再只看总分，而是逐个攻击窗口对照：应该有的战术、检测到的战术、漏掉的战术和多报的战术。")
    add_bullet(doc, "对 TRACE 和 THEIA，明确使用统一的时间偏移，避免攻击报告时间和日志时间错开。")
    add_bullet(doc, "对 THEIA，专门核对文件路径、父子进程关系、线程信息和日志中的对象字段，避免后续把真实文件或网络行为丢掉。")
    add_bullet(doc, "对 CADETS，重点审查了超大进程树为什么会出现，以及恶意进程是否被大量无关进程淹没。")

    add_heading(doc, "2.2 攻击过程识别：不只看最后有没有报出战术", 2)
    add_body(doc, "之前的优化不是简单地“多报一些战术”，而是回到攻击报告和原始日志，检查每条候选攻击过程到底对应了什么具体操作。比如，普通服务进程的启动、浏览器网络访问、文件权限改变等行为，不能因为看起来像攻击的一小部分，就被直接抬成完整攻击过程。")
    add_bullet(doc, "收紧了“前置执行”和“可疑文件落地”一类判断，尽量避免普通服务行为被当成攻击准备动作。")
    add_bullet(doc, "把敏感文件读取、外发、清理痕迹等行为改为更看重先后关系和上下文，不再只凭一次网络连接或一次文件操作下结论。")
    add_bullet(doc, "做过对照实验，分别比较是否保留关键短链、是否使用攻击先验、是否直接从候选过程判断战术，以及不同的候选数量和窗口聚合范围。")
    add_body(doc, "这些实验带来的主要认识是：前面的攻击过程质量比最后一层“选哪个战术名称”更重要。当前阶段不应该急着继续堆规则，而要先固定上游输入，再做可重复的对照。")

    add_heading(doc, "2.3 任务图和分类模型：发现不少问题其实发生在更前面", 2)
    add_body(doc, "我们也对任务图生成和恶意任务识别做了大量对照。这里的关键不是模型名字，而是日志里的进程关系被组织成什么样：如果一张图里混进了成千上万个无关进程，少量恶意操作很容易被冲淡；如果切得太碎，原本连续的攻击过程又会被拆开。")
    add_bullet(doc, "修复并核查了进程父子关系维护、线程并入进程、首条行为遗漏、文件路径读取等基础问题。")
    add_bullet(doc, "比较了只用进程行为表示、加入图整体统计信息、以及两类信息合并的不同做法。")
    add_bullet(doc, "比较了先做数据扩充再划分训练测试集，和先划分再扩充两种做法，重点避免同一张原始图的变体同时跑进训练和测试。")
    add_bullet(doc, "尝试了不同的任务图切分方式，包括按父子分支切分、是否把分支点算作上游任务的一部分，以及是否继续向上计数。")

    add_heading(doc, "3. 各数据集目前的状态", 1)
    add_status_table(doc)

    add_heading(doc, "4. 目前已经确认的几个重要结论", 1)
    add_heading(doc, "4.1 任务图切分会直接影响后面的攻击识别", 2)
    add_body(doc, "同一份日志，换一种父子关系处理或切分方式，可能会从少量很大的任务图变成更多较小的任务图。这样会改变哪些图能命中已知恶意进程，也会改变后面能够看到多少完整攻击过程。因此，比较不同版本时必须固定同一份上游任务图，不能只比较最后的战术得分。")
    add_body(doc, "尤其是 CADETS 和 E5 里都出现过“根进程下面挂了非常多子进程”的情况。这里的根进程有时只是日志生成时的占位节点，并不等于一个真实的业务任务。若直接把整棵树当作一张图，恶意行为就会被大量良性行为稀释。")

    add_heading(doc, "4.2 默认的分类判定对极少正样本很不友好", 2)
    add_body(doc, "当数万张任务图里只有十几张已知恶意图时，模型即使能把恶意图排到更靠前的位置，也可能因为默认判定门槛太保守而一张都不报。此时只看“报没报出来”会掩盖模型是否真的在进步，还需要同时看排序能力、训练集内的阈值选择和正负样本权重。")

    add_heading(doc, "4.3 最后的战术名称不是唯一瓶颈", 2)
    add_body(doc, "多次回查表明，漏报有时并不是最后一步不会命名，而是更早的时候相关任务图没有进来、文件路径没有读到、攻击过程没有被保住，或者候选过程因为数量限制被挤掉。反过来，误报也常常不是最后一步“想多了”，而是前面把普通服务活动误当成攻击链条。")

    add_heading(doc, "5. 最新 THEIA E5：接入和核验结果", 1)
    add_body(doc, "E5 的日志版本、字段结构和之前的 TC3 不完全一样，恶意名单也不是简单的一列进程编号。因此这次先没有急着追求分数，而是先把名单里每一条记录到底代表进程、文件还是网络连接核验清楚。对于文件和网络连接，我们再从日志中反查它们被哪个进程使用，并把相关进程映射到任务图上。")
    add_e5_table(doc)
    add_body(doc, "这次核验的结论很明确：文件和网络连接确实能够在日志里反查到进程，但这些进程没有带来新的恶意任务图，它们已经包含在直接标注的恶意进程范围内。因此，E5 当前效果差的主要原因不是名单漏关联，而是训练样本极少和任务图大小不均衡。")

    add_heading(doc, "5.1 重新训练后发生了什么", 2)
    add_body(doc, "为了验证旧的行为表示是否不适合 E5，我们按论文描述重新做了一次训练：让同一进程前面的行为去预测它下一次会发生什么行为。训练后，模型把恶意任务图排到前面的能力有明显提升，但使用默认判定方式时，测试集里的 4 张恶意图仍然都没有被选出来。")
    add_callout(
        doc,
        "怎么理解这个结果：",
        "不是“重新训练没用”，而是“排序已经变好，但最终挑选规则还不合适”。下一步应该在训练部分单独确定合理的正负样本权重和判定门槛，并且只在训练数据内部确定它们，不能拿测试集反复调。",
        color=DARK_BLUE,
    )

    add_heading(doc, "5.2 大任务图的现象", 2)
    add_body(doc, "E5 一共切出了 44,793 张任务图，大多数很小，但少数图特别大。最大的图有 33,818 个进程节点；一张包含已知恶意进程的图也有 3,080 个节点。它们共同的特征是：根节点下面直接或间接连着大量分支，而根节点本身没有正常的父进程信息。")
    add_body(doc, "这不一定说明日志有错，但说明当前切法对这种根节点不够细。攻击报告里的活动通常只占其中很小的一部分，整张图拿去做整体分类时，很容易被无关行为覆盖。这个问题比继续微调最后的战术规则更靠前，也更值得优先处理。")

    add_heading(doc, "6. 目前不建议做什么", 1)
    add_bullet(doc, "不建议把没有原始日志依据的战术直接塞进“已确认战术”里，只为了让分数好看。")
    add_bullet(doc, "不建议在上游任务图不断变化时横向比较最后的识别结果；这会把输入差异误认为规则差异。")
    add_bullet(doc, "不建议仅靠降低判定门槛来解决 E5 的问题。这样可能让少量恶意图被找出来，但也可能引入大量无关图，必须在训练集内先做严格对照。")
    add_bullet(doc, "不建议马上扩大所有事件类型或加入很多新规则。当前先处理根节点大分支和样本极不平衡，收益更可控。")

    add_heading(doc, "7. 建议的下一步", 1)
    add_number(doc, "先对 E5 的超大根节点做定点审查：抽取最大的几张图，回查根节点和第一层子进程在原始日志中的真实含义，确认它们是独立任务还是本来就应该放在一起。")
    add_number(doc, "在确认语义后，做一个只针对“没有正常父进程且分支极多”的根节点切分对照。保留根节点作为必要上下文，但不再让所有子分支强行共用一张超大图。")
    add_number(doc, "针对 E5 的极少正样本，做独立的训练与评估方案：固定训练测试划分，在训练部分选择正负样本权重和判定门槛，再一次性评估测试集。")
    add_number(doc, "把 E5 的任务图质量稳定后，再接入后面的攻击过程和战术识别；否则后半段看到的是被大图稀释后的输入，难以判断规则到底好不好。")
    add_number(doc, "对 TC3 三个数据集继续保持同一套逐窗口检查表。任何新优化都至少要说明：它修的是哪一种具体攻击操作、原始日志里有什么依据、对其他数据集有没有副作用。")

    add_heading(doc, "8. 总结", 1)
    add_body(doc, "目前项目不是停在“模型不行”，而是已经把几个更根本的问题拆出来了：日志与恶意名单的对齐、任务图边界、极端不平衡数据下的判定方式，以及攻击过程是否真的对应攻击报告描述。TC3 三个数据集已经积累了较完整的对照和经验；E5 已完成接入、对象关联、线程处理和重新训练验证，下一步应当优先解决超大根节点和极少正样本，而不是急着继续堆后面的战术规则。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
