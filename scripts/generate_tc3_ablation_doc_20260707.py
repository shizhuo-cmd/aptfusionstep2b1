from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:\daima\APT-Fusionstep2b1")
RAW_JSON = ROOT / "docs" / "tc3_ablation_matrix_20260706_results_raw.json"
OUT_DOCX = ROOT / "docs" / "tc3_ablation_experiments_2026-07-07_zh.docx"


TRACE_ORDER = [
    "trace_baseline_current.json",
    "trace_latefusion_off.json",
    "trace_early_concat.json",
    "trace_seq_only.json",
    "trace_fanout_gt2_include.json",
    "trace_fanout_gt3_exclude.json",
    "trace_fanout_gt3_include.json",
    "trace_connected.json",
]

CADETS_ORDER = [
    "cadets_baseline_current.json",
    "cadets_latefusion_off.json",
    "cadets_early_concat.json",
    "cadets_seq_only.json",
    "cadets_fanout_gt2_include.json",
    "cadets_fanout_gt3_exclude.json",
    "cadets_fanout_gt3_include_failed.json",
    "cadets_connected.json",
]

LABELS = {
    "trace_baseline_current.json": "Baseline（晚期融合，fanout>2，exclude）",
    "trace_latefusion_off.json": "1.1 关闭图统计晚期融合",
    "trace_early_concat.json": "1.2 早期拼接（序列+统计直接拼接）",
    "trace_seq_only.json": "1.2 仅序列表示",
    "trace_fanout_gt2_include.json": "1.3 fanout>2，include segmented",
    "trace_fanout_gt3_exclude.json": "1.3 fanout>3，exclude segmented",
    "trace_fanout_gt3_include.json": "1.3 fanout>3，include segmented",
    "trace_connected.json": "1.3 connected 切图",
    "cadets_baseline_current.json": "Baseline（晚期融合，fanout>2，exclude）",
    "cadets_latefusion_off.json": "1.1 关闭图统计晚期融合",
    "cadets_early_concat.json": "1.2 早期拼接（序列+统计直接拼接）",
    "cadets_seq_only.json": "1.2 仅序列表示",
    "cadets_fanout_gt2_include.json": "1.3 fanout>2，include segmented",
    "cadets_fanout_gt3_exclude.json": "1.3 fanout>3，exclude segmented",
    "cadets_fanout_gt3_include_failed.json": "1.3 fanout>3，include segmented（失败）",
    "cadets_connected.json": "1.3 connected 切图",
}


def load_results() -> dict:
    for encoding in ("utf-8", "utf-16", "utf-16-le"):
        try:
            return json.loads(RAW_JSON.read_text(encoding=encoding))
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("utf-8", b"", 0, 1, f"unable to decode {RAW_JSON}")


def fmt_float(value: float | int | None, digits: int = 4) -> str:
    if value is None:
        return "-"
    if isinstance(value, int):
        return str(value)
    return f"{value:.{digits}f}"


def get_pred_pos(data: dict) -> int | None:
    selected = data.get("selected_positive_task_ids")
    if isinstance(selected, list):
        return len(selected)
    module5 = data.get("module5_summary", {})
    task_count = module5.get("task_count")
    if isinstance(task_count, int):
        return task_count
    return None


def get_candidate_paths(data: dict) -> int | None:
    module5 = data.get("module5_summary", {})
    return module5.get("candidate_path_count")


def run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: int = 12, bold: bool = False):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    run_font(r, east_asia="黑体", size=10)


def set_page_layout(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def set_default_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TRACE 与 CADETS 任务图检测及下游战术分析消融实验报告")
    run_font(r, east_asia="黑体", size=16, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Ablation Study Report for Task-Graph Detection and Downstream Tactic Analysis")
    run_font(r2, east_asia="宋体", size=11)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("生成日期：2026-07-07")
    run_font(r3, size=10)


def add_section_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    if level == 1:
        r = p.add_run(text)
        run_font(r, east_asia="黑体", size=14, bold=True)
    else:
        r = p.add_run(text)
        run_font(r, east_asia="黑体", size=12, bold=True)


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    run_font(r, size=12)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(f"• {text}")
    run_font(r, size=11)


def add_experiment_design_table(doc: Document):
    add_caption(doc, "表 1  消融实验设计")
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["编号", "变量", "对照设置", "研究目的"]
    for idx, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], text, bold=True)

    rows = [
        ("1.1", "图统计晚期融合开关", "baseline vs latefusion_off", "检验 GraphSAGE 主干之外的图统计辅助分类器是否提供稳定增益。"),
        ("1.2", "节点特征组织方式", "late fusion / early concat / sequence only", "比较统计特征更适合作为独立辅助分支，还是直接并入节点主干表示。"),
        ("1.3", "任务切图策略", "fanout 阈值、include/exclude segmented、connected", "分析切图粒度对任务数量、模块2分类稳定性及下游战术识别的影响。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, item in enumerate(row):
            set_cell_text(cells[i], item, size=10)
    doc.add_paragraph()


def add_metric_note(doc: Document):
    add_body_paragraph(
        doc,
        "模块2结果使用 fit_predict 阶段的 Accuracy、Precision、Recall、F1 和 ROC-AUC；下游结果使用 confirmed_window_recall、strict_tactic_recall_macro、strict_tactic_precision_macro 与 off_window_high_risk_rate。对于 TRACE 的 1.1 与 1.2，若模块2 选中的 predicted-positive 任务与 baseline 完全一致，则 runner 直接复用 baseline 的 module3-module6 及评估结果，并在表格中显式标注。"
    )


def add_result_table(doc: Document, title: str, rows: list[dict]):
    add_caption(doc, title)
    table = doc.add_table(rows=1, cols=11)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = [
        "实验",
        "任务图数",
        "模块2 F1",
        "模块2 AUC",
        "预测正类任务数",
        "候选链条数",
        "窗口召回",
        "战术召回",
        "战术精度",
        "窗外高风险率",
        "备注",
    ]
    for idx, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], text, bold=True, size=10)

    for row in rows:
        cells = table.add_row().cells
        is_failed = "module1_summary" not in row
        if is_failed:
            error_text = row.get("failed", {}).get("error") or row.get("error", "failed")
            values = [
                LABELS[row["file"]],
                "-",
                "-",
                "-",
                "-",
                "-",
                "-",
                "-",
                "-",
                "-",
                error_text,
            ]
        else:
            note_parts = []
            if row.get("downstream_reused_from_baseline"):
                note_parts.append("后半段复用 baseline")
            if row.get("module1_summary", {}).get("task_component_split_mode") == "connected":
                note_parts.append("connected 粒度极粗")
            values = [
                LABELS[row["file"]],
                fmt_float(row["module1_summary"].get("task_count"), 0),
                fmt_float(row["module2_fit_predict_summary"]["evaluation_metrics"].get("f1")),
                fmt_float(row["module2_fit_predict_summary"]["evaluation_metrics"].get("roc_auc")),
                fmt_float(get_pred_pos(row), 0),
                fmt_float(get_candidate_paths(row), 0),
                fmt_float(row["path_reason_eval_metrics"].get("confirmed_window_recall")),
                fmt_float(row["path_reason_eval_metrics"].get("strict_tactic_recall_macro")),
                fmt_float(row["path_reason_eval_metrics"].get("strict_tactic_precision_macro")),
                fmt_float(row["path_reason_eval_metrics"].get("off_window_high_risk_rate")),
                "；".join(note_parts) if note_parts else "-",
            ]
        for i, item in enumerate(values):
            set_cell_text(cells[i], str(item), size=9)
    doc.add_paragraph()


def add_selected_task_note(doc: Document, title: str, rows: list[dict]):
    add_section_heading(doc, title, level=2)
    for row in rows:
        if "module1_summary" not in row:
            error_text = row.get("failed", {}).get("error") or row.get("error", "failed")
            add_bullet(doc, f"{LABELS[row['file']]}：失败，错误为 {error_text}。")
            continue
        selected = row.get("selected_positive_task_ids")
        if isinstance(selected, list) and selected:
            add_bullet(
                doc,
                f"{LABELS[row['file']]}：模块2 选中的 predicted-positive 任务为 {', '.join(selected)}。"
            )


def build_interpretation(doc: Document):
    add_section_heading(doc, "4 结果分析与讨论", level=1)

    add_section_heading(doc, "4.1 图统计晚期融合（1.1）", level=2)
    add_body_paragraph(
        doc,
        "在 TRACE 上，关闭图统计晚期融合后，模块2 仍然稳定选中同一组 4 张 GT 任务图，后半段结果与 baseline 完全一致，说明当前 TRACE 上的主要区分信息已经被序列表示充分吸收，图统计辅助分支并不是决定性因素。"
    )
    add_body_paragraph(
        doc,
        "在 CADETS 上，关闭晚期融合后模块2 的 F1 由 0.8745 上升到 0.9283，但 predicted-positive 任务由 10 张下降到 6 张；对应地下游 strict_tactic_recall_macro 从 0.4583 降到 0.4167。该现象表明，模块2 表面上的分类更“干净”，但它牺牲了后半段战术识别所依赖的恶意任务覆盖。"
    )

    add_section_heading(doc, "4.2 节点特征组织方式（1.2）", level=2)
    add_body_paragraph(
        doc,
        "TRACE 上，早期拼接与仅序列表示均没有改变 predicted-positive 任务集合，因此后半段结果保持不变。说明在当前 TRACE 切图条件下，统计特征既没有带来显著正增益，也没有构成明显负担。"
    )
    add_body_paragraph(
        doc,
        "CADETS 上，早期拼接与关闭晚期融合几乎表现一致：都将 predicted-positive 任务压缩到 6 张，并导致战术召回下降。相对地，仅序列表示保留了与 baseline 完全相同的 10 张任务和下游结果。这说明当前 CADETS 上，统计特征直接并入主干模型比作为独立分支更容易改变模块2 的筛选边界。"
    )

    add_section_heading(doc, "4.3 任务切图策略（1.3）", level=2)
    add_body_paragraph(
        doc,
        "TRACE 上，fanout>2 且 include segmented 与 baseline 几乎等价，只是任务图数从 1108 上升到 1209；而将阈值提高到 3 会使 predicted-positive 任务从 4 张降为 3 张，其中 exclude 版本的战术召回下降到 0.40。connected 切图则将任务图压缩到仅 14 张，直接导致 confirmed_window_recall 与 strict_tactic_recall_macro 同时归零。"
    )
    add_body_paragraph(
        doc,
        "CADETS 上，切图策略对结果影响更剧烈。fanout>2 include 与 1.1/1.2 的收缩效应一致，仍只留下 6 张 predicted-positive 任务；fanout>3 exclude 进一步降到 3 张任务，并未带来额外收益。connected 切图将任务图降到 136 张，模块2 ROC-AUC 仅 0.2963，下游 confirmed_window_recall 仅 0.25，表明过粗的图划分会破坏恶意子结构的可分性。"
    )

    add_section_heading(doc, "4.4 失败实验说明", level=2)
    add_body_paragraph(
        doc,
        "CADETS 的 fanout>3 且 include segmented 变体在本轮运行中触发 TimeoutError 并被自动移动到 `_failed_` 目录。由于该失败发生在单个实验目录内，baseline 及其他已完成产物未被覆盖。该变体是否值得重试，需要结合其更高的任务生成开销与此前 fanout>3 exclude 已经表现不佳的事实综合判断。"
    )

    add_section_heading(doc, "5 结论", level=1)
    add_body_paragraph(
        doc,
        "综合本轮消融结果，TRACE 的最优折中仍然是当前 baseline：fanout>2、exclude segmented、保留晚期融合接口但不依赖其带来额外收益。CADETS 上，如果优先追求 confirmed 窗口与战术召回，则 baseline 仍是最稳妥的选择；若只追求略低的窗外高风险率，则可以参考 latefusion_off 或 early_concat，但需要接受 predicted-positive 任务数从 10 张降到 6 张，并带来战术召回下降。"
    )
    add_body_paragraph(
        doc,
        "从论文叙述角度看，本轮结果支持一个明确结论：任务切图策略比统计特征组织方式更能决定下游战术检测的上限，而 connected 切图在 TRACE 与 CADETS 上均不适合作为当前方案的候选主线。"
    )


def create_document():
    data = load_results()
    trace_rows = [{**data[name], "file": name} for name in TRACE_ORDER if name in data]
    cadets_rows = [{**data[name], "file": name} for name in CADETS_ORDER if name in data]

    doc = Document()
    set_page_layout(doc)
    set_default_styles(doc)
    add_title_block(doc)

    add_section_heading(doc, "摘 要", level=1)
    add_body_paragraph(
        doc,
        "本文整理了 2026-07-06 在 TRACE 与 CADETS 数据集上完成的一组任务图检测与下游战术识别消融实验。实验围绕图统计晚期融合开关、节点特征组织方式以及任务切图策略三个维度展开，并统一在 augment-before-split 条件下比较模块2 的恶意任务图检测性能与下游战术检测指标。结果表明：TRACE 对 1.1 与 1.2 类改动高度稳定，真正敏感的因素是切图策略；CADETS 则对统计特征组织与切图粒度都更敏感，其中 baseline 在保持恶意任务覆盖方面最稳，而 connected 切图在两个数据集上都显著劣化。"
    )
    add_body_paragraph(
        doc,
        "关键词：任务图检测；GraphSAGE；图统计晚期融合；任务切图；TRACE；CADETS；战术识别"
    )

    add_section_heading(doc, "1 实验背景与目的", level=1)
    add_body_paragraph(
        doc,
        "本轮实验的目标是评估三类设计因素对任务图检测和下游战术分析的影响：其一，图统计特征作为独立辅助分类器时是否能提供稳定增益；其二，统计特征应当通过早期拼接还是晚期融合进入模型；其三，任务切图粒度如何影响任务数量、模块2 检测质量以及攻击窗口级战术召回。"
    )

    add_section_heading(doc, "2 实验设计", level=1)
    add_experiment_design_table(doc)
    add_metric_note(doc)

    add_section_heading(doc, "3 实验结果", level=1)
    add_result_table(doc, "表 2  TRACE 消融实验结果", trace_rows)
    add_result_table(doc, "表 3  CADETS 消融实验结果", cadets_rows)

    add_selected_task_note(doc, "3.1 CADETS 各变体选中的 predicted-positive 任务", cadets_rows)
    add_selected_task_note(doc, "3.2 TRACE 各变体选中的 predicted-positive 任务", trace_rows)

    build_interpretation(doc)
    doc.save(OUT_DOCX)


if __name__ == "__main__":
    create_document()
